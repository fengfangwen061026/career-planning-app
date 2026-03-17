"""File text extractor - extracts text from PDF and DOCX files with OCR fallback."""

import io
import logging
import os
import shutil
import zipfile
from xml.etree import ElementTree as ET
from typing import Optional, Tuple

import docx
import pdfplumber

logger = logging.getLogger(__name__)


class FileExtractor:
    """Extract text from various file formats."""

    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    MIN_TEXT_LENGTH = 100  # Minimum text length

    # Class-level flag to check if OCR is available
    _ocr_available: Optional[bool] = None

    @classmethod
    def _check_ocr_available(cls) -> bool:
        """Check if OCR dependencies are available."""
        if cls._ocr_available is not None:
            return cls._ocr_available

        try:
            # Check if tesseract is available
            tesseract_path = shutil.which("tesseract")
            if not tesseract_path:
                # Try common Windows paths
                common_paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                ]
                for path in common_paths:
                    if os.path.exists(path):
                        tesseract_path = path
                        break

            if not tesseract_path:
                cls._ocr_available = False
                logger.warning("Tesseract not found, OCR will not be available")
                return False

            # Try importing pytesseract
            import pytesseract

            # Check for Chinese language support
            try:
                langs = pytesseract.get_languages()
                if "chi_sim" not in langs:
                    logger.warning("Chinese language pack (chi_sim) not found in Tesseract")
                    # Still return True, OCR will work but may not recognize Chinese well
            except Exception:
                pass

            cls._ocr_available = True
            return True

        except ImportError:
            cls._ocr_available = False
            logger.warning("pytesseract not available, OCR will not be available")
            return False
        except Exception as e:
            cls._ocr_available = False
            logger.warning(f"OCR check failed: {e}")
            return False

    @classmethod
    def extract_text(
        cls, file_content: bytes, filename: str
    ) -> Tuple[str, list[str]]:
        """Extract text from file based on file extension.

        Args:
            file_content: Raw file bytes
            filename: File name to determine type

        Returns:
            Tuple of (extracted text, warnings list)

        Raises:
            ValueError: If file is too large or text is too short
        """
        warnings: list[str] = []

        # Check file size
        if len(file_content) > cls.MAX_FILE_SIZE:
            raise ValueError(
                f"File size exceeds {cls.MAX_FILE_SIZE / 1024 / 1024}MB limit"
            )

        # Determine file type by filename extension
        suffix = filename.lower().split(".")[-1]

        if suffix == "pdf":
            text = cls._extract_from_pdf(file_content)
        elif suffix in ("docx", "doc"):
            text = cls._extract_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: .{suffix}")

        # If text is too short, try OCR as fallback
        if len(text.strip()) < cls.MIN_TEXT_LENGTH:
            ocr_text = cls._extract_with_ocr(file_content, filename)
            if ocr_text and len(ocr_text.strip()) >= cls.MIN_TEXT_LENGTH:
                text = ocr_text
                warnings.append(
                    "简历为图片格式，已通过 OCR 识别，建议使用文字版简历以获得更准确的解析结果"
                )

        # Final check
        if len(text.strip()) < cls.MIN_TEXT_LENGTH:
            raise ValueError(
                f"Extracted text is too short ({len(text)} chars). "
                f"Minimum required: {cls.MIN_TEXT_LENGTH} chars. "
                f"Supported formats: text-based PDF, DOCX, or image-based PDF with tesseract OCR."
            )

        return text.strip(), warnings

    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF using pdfplumber.

        Args:
            file_content: PDF file bytes

        Returns:
            Extracted text
        """
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            pages: list[str] = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n".join(pages)

    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX using python-docx.

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text
        """
        doc = docx.Document(io.BytesIO(file_content))
        parts: list[str] = []

        for para in doc.paragraphs:
            try:
                text = para.text.strip()
            except Exception as exc:
                logger.warning("Skipping unreadable DOCX paragraph: %s", exc)
                continue
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_texts: list[str] = []
                for cell in row.cells:
                    try:
                        cell_text = cell.text.strip()
                    except Exception as exc:
                        logger.warning("Skipping unreadable DOCX table cell: %s", exc)
                        continue
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append("\t".join(row_texts))

        textbox_text = FileExtractor._extract_docx_textboxes(file_content)
        if textbox_text:
            parts.extend(textbox_text)

        return "\n".join(parts)

    @staticmethod
    def _extract_docx_textboxes(file_content: bytes) -> list[str]:
        """Extract text from DOCX text boxes and shapes.

        Some resume templates store nearly all visible text inside `w:txbxContent`,
        which `python-docx` does not expose via `doc.paragraphs`.
        """
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as archive:
                xml = archive.read("word/document.xml")
        except Exception as exc:
            logger.warning("Failed to read DOCX XML for textbox extraction: %s", exc)
            return []

        try:
            root = ET.fromstring(xml)
        except Exception as exc:
            logger.warning("Failed to parse DOCX XML for textbox extraction: %s", exc)
            return []

        textbox_parts: list[str] = []
        for textbox in root.findall(".//w:txbxContent", namespace):
            for para in textbox.findall(".//w:p", namespace):
                texts: list[str] = []
                for text_node in para.findall(".//w:t", namespace):
                    text = (text_node.text or "").strip()
                    if text:
                        texts.append(text)
                if texts:
                    textbox_parts.append("".join(texts))

        return textbox_parts

    @classmethod
    def _extract_with_ocr(cls, file_content: bytes, filename: str) -> str:
        """Extract text using OCR when normal extraction fails.

        Args:
            file_content: File bytes
            filename: Original filename

        Returns:
            OCR extracted text, or empty string if OCR fails
        """
        if not cls._check_ocr_available():
            return ""

        try:
            import pytesseract
            from pdf2image import convert_from_bytes
            from PIL import Image

            suffix = filename.lower().split(".")[-1]

            # Configure tesseract (Windows may need explicit path)
            tesseract_path = shutil.which("tesseract")
            if not tesseract_path:
                common_paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                ]
                for path in common_paths:
                    if os.path.exists(path):
                        tesseract_path = path
                        break

            if tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = tesseract_path

            text_parts: list[str] = []

            if suffix == "pdf":
                # Convert PDF to images
                images = convert_from_bytes(file_content)
                for image in images:
                    text = pytesseract.image_to_string(
                        image, lang="chi_sim+eng"
                    )
                    if text:
                        text_parts.append(text)
            elif suffix in ("docx", "doc"):
                # For DOCX with images, we can't easily OCR
                # Return empty to indicate failure
                logger.info("OCR not supported for DOCX with images")
                return ""

            return "\n".join(text_parts)

        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
            return ""


def extract_text(file_content: bytes, filename: str) -> Tuple[str, list[str]]:
    """Convenience function for text extraction.

    Args:
        file_content: Raw file bytes
        filename: File name to determine type

    Returns:
        Tuple of (extracted text, warnings list)
    """
    return FileExtractor.extract_text(file_content, filename)
