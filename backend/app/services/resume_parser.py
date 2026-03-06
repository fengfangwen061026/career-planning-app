"""Resume parser service - parses resumes and extracts structured information.

Pipeline:
  1. 文本提取 - PDF(PyMuPDF) / DOCX(python-docx)
  2. LLM 结构化抽取 - 调用 resume_parse prompt
  3. 技能归一化 - 同义词词表硬匹配
  4. 学生画像组装 - 完整度评分 + 缺失建议
"""

import logging
from pathlib import Path
from typing import Any
from uuid import UUID

import fitz  # PyMuPDF
from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.llm_provider import llm
from app.ai.prompts.resume_parse import build_resume_parse_prompt
from app.models.student import Resume, Student
from app.utils.skill_normalizer import normalize

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# 1. 文本提取
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    """使用 PyMuPDF 从 PDF 提取文本。"""
    doc = fitz.open(file_path)
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages).strip()


def extract_text_from_docx(file_path: str) -> str:
    """使用 python-docx 从 DOCX 提取文本。"""
    doc = Document(file_path)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip()


def extract_text_from_file(file_path: str) -> str:
    """根据文件类型分发提取。"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# 2. LLM 结构化抽取
# ---------------------------------------------------------------------------

async def parse_with_ai(raw_text: str) -> dict[str, Any]:
    """调用 LLM 解析简历文本为结构化 JSON。"""
    messages = build_resume_parse_prompt(raw_text)
    system_prompt = messages[0]["content"]
    user_prompt = messages[1]["content"]

    parsed = await llm.generate_json(
        prompt=user_prompt,
        system_prompt=system_prompt,
        temperature=0.2,
        max_retries=3,
    )
    logger.info("LLM resume parsing done, confidence=%.2f",
                parsed.get("_meta", {}).get("parse_confidence", 0))
    return parsed


# ---------------------------------------------------------------------------
# 3. 技能归一化
# ---------------------------------------------------------------------------

def normalize_parsed_skills(
    parsed_data: dict[str, Any],
) -> tuple[dict[str, Any], list[dict[str, str]]]:
    """对解析结果中的技能做归一化，返回 (更新后的数据, 归一化日志)。"""
    normalization_log: list[dict[str, str]] = []
    skills = parsed_data.get("skills", [])

    for skill_entry in skills:
        original = skill_entry.get("name", "")
        normalized = normalize(original)
        if normalized != original:
            normalization_log.append({
                "original": original,
                "normalized": normalized,
            })
            skill_entry["name"] = normalized

    # 也归一化 project_experience 中的 tech_stack
    for proj in parsed_data.get("project_experience", []):
        tech_stack = proj.get("tech_stack", [])
        normalized_stack: list[str] = []
        for tech in tech_stack:
            norm = normalize(tech)
            if norm != tech:
                normalization_log.append({
                    "original": tech,
                    "normalized": norm,
                })
            normalized_stack.append(norm)
        proj["tech_stack"] = normalized_stack

    return parsed_data, normalization_log


# ---------------------------------------------------------------------------
# 4. 完整度评分 & 缺失建议
# ---------------------------------------------------------------------------

# 各维度权重
_DIMENSION_WEIGHTS: dict[str, float] = {
    "basic_info": 0.10,
    "education": 0.20,
    "work_experience": 0.20,
    "project_experience": 0.15,
    "skills": 0.20,
    "certificates": 0.05,
    "awards": 0.05,
    "soft_skills": 0.05,
}


def _dimension_fill_rate(parsed_data: dict[str, Any], dim: str) -> float:
    """计算单个维度的填充率 (0-1)。"""
    val = parsed_data.get(dim)
    if val is None:
        return 0.0

    if dim == "basic_info":
        if not isinstance(val, dict):
            return 0.0
        key_fields = ["name", "phone", "email", "education", "location", "job_intention"]
        filled = sum(1 for k in key_fields if val.get(k))
        return filled / len(key_fields)

    if dim == "soft_skills":
        if not isinstance(val, dict):
            return 0.0
        scored = sum(1 for v in val.values() if isinstance(v, dict) and v.get("score", 0) > 1)
        return min(scored / 6, 1.0)

    # 列表类型维度
    if isinstance(val, list):
        if not val:
            return 0.0
        # 有内容就按条目数给分（1 条 0.5，2 条 0.7，3+ 条 1.0）
        n = len(val)
        if n >= 3:
            return 1.0
        elif n == 2:
            return 0.7
        else:
            return 0.5

    return 0.0


def compute_completeness_score(parsed_data: dict[str, Any]) -> float:
    """计算画像完整度评分 (0-1)，各维度加权。"""
    score = 0.0
    for dim, weight in _DIMENSION_WEIGHTS.items():
        rate = _dimension_fill_rate(parsed_data, dim)
        score += rate * weight
    return round(score, 3)


def generate_missing_suggestions(parsed_data: dict[str, Any]) -> list[str]:
    """根据解析结果生成缺失项建议。"""
    suggestions: list[str] = []

    # 基本信息
    basic = parsed_data.get("basic_info", {})
    if not basic.get("email"):
        suggestions.append("建议补充邮箱地址")
    if not basic.get("phone"):
        suggestions.append("建议补充手机号码")
    if not basic.get("job_intention"):
        suggestions.append("建议补充求职意向")

    # 教育
    education = parsed_data.get("education", [])
    if not education:
        suggestions.append("缺少教育经历，请补充学校、专业、学历等信息")
    else:
        for edu in education:
            if not edu.get("gpa"):
                suggestions.append("建议补充 GPA/成绩排名等量化指标")
                break

    # 工作/实习经验
    work = parsed_data.get("work_experience", [])
    if not work:
        suggestions.append("缺少工作/实习经历")
    else:
        has_achievement = any(
            exp.get("achievements") for exp in work
        )
        if not has_achievement:
            suggestions.append("建议在工作经历中补充量化成果（如提升了 XX%、负责了 XX 规模的项目）")

    # 项目经验
    projects = parsed_data.get("project_experience", [])
    if not projects:
        suggestions.append("缺少项目经历，建议补充个人项目、课程项目或开源贡献")
    else:
        has_tech = any(proj.get("tech_stack") for proj in projects)
        if not has_tech:
            suggestions.append("建议在项目经历中注明技术栈")
        has_achievement = any(proj.get("achievements") for proj in projects)
        if not has_achievement:
            suggestions.append("建议在项目经历中补充量化成果")

    # 技能
    skills = parsed_data.get("skills", [])
    if not skills:
        suggestions.append("缺少技能列表，请补充专业技能")

    # 证书
    certs = parsed_data.get("certificates", [])
    if not certs:
        suggestions.append("建议补充相关证书（如语言证书、专业认证等）")

    # _meta 中标记的缺失
    meta = parsed_data.get("_meta", {})
    for field in meta.get("missing_fields", []):
        suggestions.append(f"简历中未找到：{field}")

    return suggestions


# ---------------------------------------------------------------------------
# 5. 主流程
# ---------------------------------------------------------------------------

async def parse_resume(
    resume_id: UUID,
    db: AsyncSession,
) -> dict[str, Any]:
    """完整的简历解析流水线。

    1. 读取 Resume 记录 → 提取文本
    2. LLM 结构化抽取
    3. 技能归一化
    4. 计算完整度 + 缺失建议
    5. 更新数据库

    Returns:
        {
            "resume": Resume ORM object,
            "parsed_data": dict,
            "completeness_score": float,
            "missing_suggestions": list[str],
            "normalization_log": list[dict],
        }
    """
    # 获取 Resume
    resume = await db.get(Resume, resume_id)
    if not resume:
        raise ValueError(f"Resume {resume_id} not found")

    if not resume.file_path:
        raise ValueError(f"Resume {resume_id} has no file_path")

    # Step 1: 文本提取
    logger.info("Extracting text from %s", resume.file_path)
    raw_text = extract_text_from_file(resume.file_path)
    if not raw_text.strip():
        raise ValueError("Extracted text is empty - file may be an image-only PDF")
    resume.raw_text = raw_text

    # Step 2: LLM 结构化抽取
    logger.info("Parsing resume with LLM (resume_id=%s)", resume_id)
    parsed_data = await parse_with_ai(raw_text)

    # Step 3: 技能归一化
    parsed_data, normalization_log = normalize_parsed_skills(parsed_data)
    if normalization_log:
        logger.info("Normalized %d skills: %s",
                     len(normalization_log),
                     ", ".join(f"{e['original']}→{e['normalized']}" for e in normalization_log))

    # Step 4: 完整度评分 & 缺失建议
    completeness_score = compute_completeness_score(parsed_data)
    missing_suggestions = generate_missing_suggestions(parsed_data)

    # Step 5: 保存到数据库
    resume.parsed_json = parsed_data
    await db.flush()
    await db.refresh(resume)

    logger.info("Resume %s parsed successfully (completeness=%.2f)", resume_id, completeness_score)

    return {
        "resume": resume,
        "parsed_data": parsed_data,
        "completeness_score": completeness_score,
        "missing_suggestions": missing_suggestions,
        "normalization_log": normalization_log,
    }


async def update_student_basic_info(
    student_id: UUID,
    parsed_data: dict[str, Any],
    db: AsyncSession,
) -> None:
    """从解析结果更新 Student 表的基本信息。"""
    student = await db.get(Student, student_id)
    if not student:
        return

    basic = parsed_data.get("basic_info", {})

    if basic.get("name") and not student.name:
        student.name = basic["name"]
    if basic.get("phone") and not student.phone:
        student.phone = basic["phone"]
    if basic.get("gender") and not student.gender:
        student.gender = basic["gender"]
    if basic.get("location") and not student.location:
        student.location = basic["location"]
    if basic.get("hometown") and not student.hometown:
        student.hometown = basic["hometown"]
    if basic.get("job_intention") and not student.job_intention:
        student.job_intention = basic["job_intention"]

    await db.flush()
