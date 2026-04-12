"""Career report generation, normalization, and export services."""

from __future__ import annotations

import asyncio
import html
import json
import logging
import re
import subprocess
import sys
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from uuid import UUID

from sqlalchemy import desc, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.llm_provider import llm
from app.models.report import CareerReport, ReportVersion
from app.models.student import Student, StudentProfile
from app.prompts.report_generation import (
    REPORT_POLISH_SYSTEM_PROMPT,
    REPORT_POLISH_USER_TEMPLATE,
    REPORT_SYSTEM_PROMPT,
    CHAPTER_1_PROMPT,
    CHAPTER_2_PROMPT,
    CHAPTER_3_PROMPT,
    CHAPTER_4_PROMPT,
    CHAPTER_5_PROMPT,
)
from app.services.graph import find_path_with_student_profile
from app.services.matching import match_student_job, recommend_jobs
from app.services.student_profile import normalize_student_profile_json

logger = logging.getLogger(__name__)

EXPORT_DIR = Path(__file__).resolve().parents[2] / "static" / "exports"
REPORT_TEMPLATE_VERSION = "2.0"
DEFAULT_REPORT_TITLE = "职业发展报告"
REPORT_CHAPTERS = [
    {"chapter_id": 1, "title": "一、个人优势总结"},
    {"chapter_id": 2, "title": "二、目标岗位分析"},
    {"chapter_id": 3, "title": "三、差距与行动计划"},
    {"chapter_id": 4, "title": "四、职业路径规划"},
    {"chapter_id": 5, "title": "五、评估周期"},
]
DIMENSION_META = [
    ("basic", "基础"),
    ("skill", "技能"),
    ("competency", "素养"),
    ("potential", "潜力"),
]
GAP_ITEM_LABELS = {
    "major_relevance": "专业与岗位方向匹配度",
}


def _normalize_score(value: Any) -> int:
    try:
        score = float(value or 0)
    except (TypeError, ValueError):
        return 0
    if score <= 1:
        score *= 100
    return max(0, min(100, round(score)))


def _safe_text(value: Any, default: str = "") -> str:
    text = str(value or "").strip()
    return text or default


def _chapter_title(chapter_id: int) -> str:
    for item in REPORT_CHAPTERS:
        if item["chapter_id"] == chapter_id:
            return item["title"]
    return f"第 {chapter_id} 章"


def _escape_html(text: Any) -> str:
    return html.escape(str(text or ""), quote=True)


def _join_non_empty(parts: list[str], separator: str = "，") -> str:
    return separator.join([part for part in parts if part])


def _clean_paragraph(text: str) -> str:
    text = " ".join(str(text or "").split()).replace("锟?", "").strip()
    if text and text[-1] not in "。！？!?":
        text += "。"
    return text


def _normalize_gap_item(raw_item: Any) -> str:
    text = _safe_text(raw_item, "能力差距")
    if text in GAP_ITEM_LABELS:
        return GAP_ITEM_LABELS[text]
    if text.startswith("必备技能:"):
        return f"必备技能：{text.split(':', 1)[-1].strip()}"
    if text.startswith("优选技能:"):
        return f"加分技能：{text.split(':', 1)[-1].strip()}"
    if text.startswith("职业素养:"):
        return f"职业素养：{text.split(':', 1)[-1].strip()}"
    if text.startswith("发展潜力:"):
        return f"发展潜力：{text.split(':', 1)[-1].strip()}"
    return text


def _build_gap_description(item_name: str, current_level: str, required_level: str) -> str:
    if item_name == GAP_ITEM_LABELS["major_relevance"]:
        return "当前专业背景与目标岗位方向的直接相关性偏弱，需要补足更贴近岗位的项目或实习证据"
    if current_level and required_level:
        return f"{current_level} -> {required_level}"
    if current_level:
        return f"当前水平：{current_level}"
    if required_level:
        return f"目标要求：{required_level}"
    return "当前能力与目标岗位要求存在差距"


def _build_gap_action(item_name: str, suggestion: str) -> str:
    if suggestion and suggestion != "建议针对该项短板制定补齐计划":
        return suggestion
    if item_name == GAP_ITEM_LABELS["major_relevance"]:
        return "补充与目标岗位更相关的课程、项目或实习经历，并在简历中明确说明转向理由与可迁移能力。"
    if item_name.startswith("必备技能："):
        skill_name = item_name.split("：", 1)[-1]
        return f"围绕 {skill_name} 完成 1 个可展示项目，并在简历中写清使用场景、技术细节和结果。"
    if item_name.startswith("加分技能："):
        skill_name = item_name.split("：", 1)[-1]
        return f"补充 {skill_name} 的基础实践，把它沉淀成面试中的加分项，而不是停留在概念层面。"
    if item_name.startswith("职业素养："):
        label = item_name.split("：", 1)[-1]
        return f"通过团队协作、复盘记录或项目推进过程，补足“{label}”这类软性能力证据。"
    if item_name.startswith("发展潜力："):
        label = item_name.split("：", 1)[-1]
        return f"主动发起新项目、比赛或作品集建设，用连续输出证明“{label}”的成长潜力。"
    return "建议围绕该短板补充一段可验证的项目、课程或实习经历，并同步更新简历表达。"


def _build_action_resources(item_name: str) -> list[str]:
    if item_name == GAP_ITEM_LABELS["major_relevance"]:
        return ["补充对口课程/项目案例", "准备一段能说明转岗理由的自我介绍"]
    if item_name.startswith("必备技能："):
        skill_name = item_name.split("：", 1)[-1]
        return [f"{skill_name} 官方文档", f"{skill_name} 实战项目模板"]
    if item_name.startswith("加分技能："):
        skill_name = item_name.split("：", 1)[-1]
        return [f"{skill_name} 入门课程", f"{skill_name} 进阶案例拆解"]
    if item_name.startswith("职业素养："):
        return ["STAR 法则复盘模板", "项目复盘记录"]
    if item_name.startswith("发展潜力："):
        return ["作品集整理清单", "比赛/开源项目计划表"]
    return ["岗位 JD 复盘", "项目成果量化模板"]


def _score_band_text(score: int) -> str:
    if score >= 80:
        return "具备较强切入基础"
    if score >= 60:
        return "具备一定切入基础"
    return "仍处在重点准备阶段"


def _normalize_action_items(items: Any) -> list[dict[str, Any]]:
    if isinstance(items, dict):
        raw_items = []
        for group_name, default_priority in (("short_term", "必须补齐"), ("medium_term", "建议提升")):
            for item in items.get(group_name) or []:
                if isinstance(item, dict):
                    raw_items.append({"priority": default_priority, **item})
    else:
        raw_items = list(items or [])

    normalized = []
    for item in raw_items:
        if not isinstance(item, dict):
            continue
        item_name = _normalize_gap_item(item.get("item") or item.get("gap_item"))
        current_level = _safe_text(item.get("current_level"))
        required_level = _safe_text(item.get("required_level"))
        gap_desc = _safe_text(item.get("gap_desc"))
        if item_name == GAP_ITEM_LABELS["major_relevance"] or not gap_desc:
            gap_desc = _build_gap_description(item_name, current_level, required_level)
        priority = _safe_text(item.get("priority"), "持续巩固")
        score_impact = int(item.get("score_impact") or 0)
        if score_impact == 0:
            if priority == "必须补齐":
                score_impact = -12
            elif priority == "建议提升":
                score_impact = -6
            else:
                score_impact = -3
        timeline = _safe_text(item.get("timeline"))
        if not timeline:
            if priority == "必须补齐":
                timeline = "3周内"
            elif priority == "建议提升":
                timeline = "3个月内"
            else:
                timeline = "1个季度内"
        resources = [
            _safe_text(resource)
            for resource in (item.get("resources") or [])
            if _safe_text(resource)
        ] or _build_action_resources(item_name)
        normalized.append(
            {
                "priority": priority,
                "item": item_name,
                "gap_desc": gap_desc,
                "score_impact": score_impact,
                "action": _build_gap_action(item_name, _safe_text(item.get("action") or item.get("suggestion"))),
                "timeline": timeline,
                "resources": resources[:2],
            }
        )
    return normalized


def _group_action_items(items: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped = {"short_term": [], "medium_term": []}
    for item in items:
        if not isinstance(item, dict):
            continue
        score_impact = abs(int(item.get("score_impact") or 0))
        priority = _safe_text(item.get("priority"))
        target_group = "short_term" if priority == "必须补齐" or score_impact >= 8 else "medium_term"
        grouped[target_group].append(
            {
                "priority": _safe_text(item.get("priority"), "建议提升"),
                "item": _safe_text(item.get("item"), "能力差距"),
                "gap_desc": _safe_text(item.get("gap_desc")),
                "score_impact": int(item.get("score_impact") or 0),
                "action": _safe_text(item.get("action")),
                "timeline": _safe_text(item.get("timeline")),
                "resources": [
                    _safe_text(resource)
                    for resource in (item.get("resources") or [])
                    if _safe_text(resource)
                ][:2],
            }
        )

    grouped["short_term"] = grouped["short_term"][:4]
    grouped["medium_term"] = grouped["medium_term"][:4]
    return grouped


def _build_review_checkpoints(actions: list[dict[str, Any]], baseline_score: int) -> dict[str, list[dict[str, Any]]]:
    first_action = actions[0] if actions else {
        "item": "核心岗位技能",
        "action": "补齐目标岗位的核心技能，并形成 1 个可展示成果。",
    }
    medium_action = next(
        (item for item in actions if _safe_text(item.get("priority")) != "必须补齐"),
        actions[1] if len(actions) > 1 else first_action,
    )
    month_three_target = min(98, max(baseline_score + 5, baseline_score))
    month_six_target = min(100, max(baseline_score + 10, month_three_target + 3))
    return {
        "review_checkpoints": [
            {
                "month": 1,
                "goal": f"完成{_safe_text(first_action.get('item'), '核心技能')}的首轮补齐",
                "kpi": f"产出1个与{_safe_text(first_action.get('item'), '目标岗位技能')}相关的练习或项目成果",
                "action": _safe_text(first_action.get("action"), "完成一轮专项补齐，并整理可展示证据"),
            },
            {
                "month": 3,
                "goal": f"综合分提升至{month_three_target}分",
                "kpi": f"重新运行匹配系统，综合分达到{month_three_target}分",
                "action": "更新简历并重新上传，验证差距项是否显著减少",
            },
            {
                "month": 6,
                "goal": f"完成{_safe_text(medium_action.get('item'), '中期能力')}等中期目标",
                "kpi": f"完成2个岗位相关作品或获得1段相关实习/项目经历，综合分接近{month_six_target}分",
                "action": _safe_text(medium_action.get("action"), "把中期目标转化为真实经历，并重新生成完整报告"),
            },
        ]
    }


def _normalize_review_checkpoints(
    checkpoints: Any,
    actions: list[dict[str, Any]],
    baseline_score: int,
) -> dict[str, list[dict[str, Any]]]:
    if isinstance(checkpoints, dict) and isinstance(checkpoints.get("review_checkpoints"), list):
        normalized = []
        for item in checkpoints.get("review_checkpoints") or []:
            if not isinstance(item, dict):
                continue
            try:
                month = int(item.get("month") or 0)
            except (TypeError, ValueError):
                month = 0
            normalized.append(
                {
                    "month": month,
                    "goal": _safe_text(item.get("goal"), "完成阶段性目标"),
                    "kpi": _safe_text(item.get("kpi"), "完成可验证成果并更新匹配结果"),
                    "action": _safe_text(item.get("action"), "更新简历并重新复盘"),
                }
            )
        if normalized:
            return {"review_checkpoints": normalized[:3]}
    return _build_review_checkpoints(actions, baseline_score)


def _normalize_paths(paths: Any, target_role: str, actions: list[dict[str, Any]]) -> dict[str, Any]:
    source = dict(paths or {})
    primary_path = [
        item for item in source.get("primary_path") or [] if isinstance(item, dict)
    ]
    alt_paths = [
        {"title": _safe_text(item.get("title"), "相关岗位"), "skill_overlap": _normalize_score(item.get("skill_overlap"))}
        for item in (source.get("alt_paths") or [])
        if isinstance(item, dict)
    ]
    if target_role and (
        not primary_path
        or target_role not in _safe_text(primary_path[0].get("title"))
    ):
        return {
            "primary_path": [
                {
                    "stage": "现在",
                    "title": f"{target_role} 准备期",
                    "condition": _safe_text(
                        actions[0]["action"] if actions else None,
                        "补齐岗位核心能力和求职材料，先把第一段可验证证据做出来。",
                    ),
                    "is_current": True,
                },
                {
                    "stage": "1-2年",
                    "title": f"{target_role}（初级）",
                    "condition": "完成相关实习或项目，形成稳定的岗位胜任力。",
                    "is_current": False,
                },
                {
                    "stage": "3-5年",
                    "title": f"{target_role}（进阶）",
                    "condition": "持续积累复杂项目经验和跨团队协作能力，逐步承担更大范围的结果责任。",
                    "is_current": False,
                },
            ],
            "alt_paths": alt_paths[:3],
        }
    return {"primary_path": primary_path, "alt_paths": alt_paths[:3]}


def normalize_report_content(content_json: dict[str, Any] | None) -> dict[str, Any]:
    """Normalize legacy and current report content into one shape."""
    content = dict(content_json or {})
    chapters = []
    raw_chapters = content.get("chapters") or []
    for index, item in enumerate(raw_chapters, start=1):
        if not isinstance(item, dict):
            continue
        if "text" in item:
            text = _clean_paragraph(str(item.get("text") or "")) if item.get("text") else ""
            data = item.get("data")
        else:
            sections = item.get("sections") or []
            section_texts = []
            for section in sections:
                if not isinstance(section, dict):
                    continue
                title = _safe_text(section.get("title"))
                content_text = _safe_text(section.get("content"))
                section_texts.append(f"{title}：{content_text}" if title and content_text else content_text)
            text = _clean_paragraph(" ".join([s for s in section_texts if s])) if section_texts else ""
            data = item.get("data")
        chapters.append(
            {
                "chapter_id": item.get("chapter_id") or index,
                "title": item.get("title") or _chapter_title(index),
                "text": text,
                "data": data,
                "status": item.get("status") or ("done" if text else "pending"),
            }
        )
    for chapter in REPORT_CHAPTERS:
        if not any(item["chapter_id"] == chapter["chapter_id"] for item in chapters):
            chapters.append(
                {
                    "chapter_id": chapter["chapter_id"],
                    "title": chapter["title"],
                    "text": "",
                    "data": None,
                    "status": "pending",
                }
            )
    chapters.sort(key=lambda item: item["chapter_id"])
    target_job = dict(content.get("target_job") or {})
    baseline_score = _normalize_score(
        target_job.get("match_score") or target_job.get("overall_score")
    )
    actions = _normalize_action_items(content.get("actions") or content.get("action_groups") or [])
    action_groups = _group_action_items(actions)
    review_checkpoints = _normalize_review_checkpoints(
        content.get("review_checkpoints"),
        actions,
        baseline_score,
    )
    target_role = _safe_text(
        target_job.get("role_name") or target_job.get("role") or target_job.get("title")
    )
    paths = _normalize_paths(content.get("paths"), target_role, actions)
    for chapter in chapters:
        if chapter["chapter_id"] == 3:
            chapter_data = chapter.get("data")
            if isinstance(chapter_data, dict) and (
                isinstance(chapter_data.get("short_term"), list)
                or isinstance(chapter_data.get("medium_term"), list)
            ):
                chapter["data"] = _group_action_items(_normalize_action_items(chapter_data))
            elif isinstance(chapter_data, list):
                chapter["data"] = _group_action_items(_normalize_action_items(chapter_data))
            else:
                chapter["data"] = action_groups
        elif chapter["chapter_id"] == 4:
            chapter["data"] = paths
        elif chapter["chapter_id"] == 5:
            chapter_data = chapter.get("data")
            chapter["data"] = _normalize_review_checkpoints(chapter_data, actions, baseline_score)
    return {
        "title": content.get("title") or DEFAULT_REPORT_TITLE,
        "summary": _clean_paragraph(str(content.get("summary") or "")) if content.get("summary") else "",
        "target_job": target_job,
        "dimensions": list(content.get("dimensions") or []),
        "actions": actions,
        "action_groups": action_groups,
        "paths": paths,
        "review_checkpoints": review_checkpoints,
        "chapters": chapters,
        "metadata": dict(content.get("metadata") or {}),
    }


def serialize_career_report(report: CareerReport) -> dict[str, Any]:
    content = normalize_report_content(report.content_json or {})
    return {
        "id": report.id,
        "student_id": report.student_id,
        "title": content.get("title") or DEFAULT_REPORT_TITLE,
        "summary": report.summary or content.get("summary") or "",
        "recommendations": report.recommendations or [],
        "suggested_jobs": None,
        "skill_gaps": content.get("actions"),
        "career_path": (content.get("paths") or {}).get("primary_path"),
        "status": report.status,
        "version": report.version or "1.0",
        "content_json": content,
        "pdf_path": report.pdf_path,
        "docx_path": report.docx_path,
        "created_at": report.created_at,
        "updated_at": report.updated_at,
    }


def _extract_skills(profile_json: dict[str, Any], limit: int = 6) -> list[str]:
    result = []
    for item in profile_json.get("skills") or []:
        if isinstance(item, dict):
            name = item.get("name") or item.get("skill_name")
        else:
            name = item
        if name:
            result.append(str(name))
    return result[:limit]


def _extract_projects(profile_json: dict[str, Any], limit: int = 3) -> list[str]:
    result = []
    experience = profile_json.get("experience") or {}
    for item in experience.get("projects") or profile_json.get("experiences") or []:
        if isinstance(item, dict):
            name = item.get("project_name") or item.get("name") or item.get("title")
            if name:
                result.append(str(name))
    return result[:limit]


def _extract_internships(profile_json: dict[str, Any], limit: int = 3) -> list[str]:
    result = []
    experience = profile_json.get("experience") or {}
    for item in experience.get("work") or profile_json.get("experiences") or []:
        if not isinstance(item, dict):
            continue
        company = _safe_text(item.get("company"))
        title = _safe_text(item.get("title") or item.get("position"))
        if company or title:
            result.append(_join_non_empty([company, title], " / "))
    return result[:limit]

def _extract_certificates(profile_json: dict[str, Any], limit: int = 3) -> list[str]:
    result = []
    seen: set[str] = set()
    sources = list(profile_json.get("certificates") or [])
    sources.extend(profile_json.get("certificate_names") or [])
    for item in sources:
        if isinstance(item, dict):
            name = item.get("name") or item.get("title") or item.get("certificate")
        else:
            name = item
        cleaned = _safe_text(name)
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        result.append(cleaned)
        if len(result) >= limit:
            break
    return result


def _extract_job_info(match_result: dict[str, Any]) -> dict[str, Any]:
    scores_json = dict(match_result.get("scores_json") or {})
    job_info = dict(scores_json.get("job_info") or {})
    role_name = (
        job_info.get("role")
        or job_info.get("title")
        or match_result.get("role_name")
        or match_result.get("job_title")
        or "目标岗位"
    )
    return {
        **job_info,
        "role_name": role_name,
        "match_score": _normalize_score(match_result.get("total_score") or scores_json.get("total_score")),
    }


def _extract_dimensions(match_result: dict[str, Any]) -> list[dict[str, Any]]:
    scores_json = dict(match_result.get("scores_json") or {})
    dimensions = []
    for key, label in DIMENSION_META:
        payload = dict(scores_json.get(key) or {})
        reason = _safe_text(payload.get("reason"))
        if not reason and isinstance(payload.get("items"), list) and payload["items"]:
            first_item = payload["items"][0]
            if isinstance(first_item, dict):
                reason = _safe_text(first_item.get("evidence") or first_item.get("dimension"))
        dimensions.append(
            {
                "key": key,
                "label": label,
                "score": _normalize_score(payload.get("score")),
                "reason": reason or "可继续补强这一维度的证据表达。",
            }
        )
    return dimensions


def _build_actions(match_result: dict[str, Any]) -> list[dict[str, Any]]:
    actions = []
    for item in match_result.get("gaps_json") or []:
        if not isinstance(item, dict):
            continue
        item_name = _normalize_gap_item(item.get("gap_item") or item.get("item"))
        current_level = _safe_text(item.get("current_level"))
        required_level = _safe_text(item.get("required_level"))
        raw_priority = str(item.get("priority") or "").lower()
        if raw_priority == "high":
            priority = "必须补齐"
            score_impact = -12
            timeline = "3周内"
        elif raw_priority == "medium":
            priority = "建议提升"
            score_impact = -6
            timeline = "3个月内"
        else:
            priority = "持续巩固"
            score_impact = -3
            timeline = "6个月内"
        actions.append(
            {
                "priority": priority,
                "item": item_name,
                "gap_desc": _build_gap_description(item_name, current_level, required_level),
                "score_impact": score_impact,
                "action": _build_gap_action(item_name, _safe_text(item.get("suggestion"))),
                "timeline": timeline,
                "resources": _build_action_resources(item_name),
            }
        )
    if actions:
        return actions[:5]
    return [
        {
            "priority": "持续巩固",
            "item": "强化项目表达",
            "gap_desc": "当前缺少明确阻塞项，但需要更强的求职证据",
            "score_impact": -3,
            "action": "补充项目量化成果、技术选型理由和个人贡献，持续优化简历表达。",
            "timeline": "6个月内",
            "resources": ["项目量化成果模板", "简历 STAR 表达清单"],
        }
    ]


def _build_paths(
    student_profile: dict[str, Any],
    target_role: str,
    career_path: dict[str, Any] | None,
    related_jobs: list[dict[str, Any]],
    actions: list[dict[str, Any]],
) -> dict[str, Any]:
    action_plan = list((career_path or {}).get("action_plan") or [])
    primary_path = [
        {
            "stage": "现在",
            "title": f"{target_role} 准备期",
            "condition": _safe_text(
                actions[0]["action"] if actions else None,
                "补齐岗位核心能力和求职材料，先把第一段可验证证据做出来。",
            ),
            "is_current": True,
        },
        {
            "stage": "1-2年",
            "title": f"{target_role}（初级）",
            "condition": _safe_text(
                action_plan[0].get("action") if action_plan and isinstance(action_plan[0], dict) else None,
                "完成相关实习或项目，形成稳定的岗位胜任力。",
            ),
            "is_current": False,
        },
        {
            "stage": "3-5年",
            "title": f"{target_role}（进阶）",
            "condition": _safe_text(
                action_plan[1].get("action") if len(action_plan) > 1 and isinstance(action_plan[1], dict) else None,
                "持续积累复杂项目经验和跨团队协作能力，逐步承担更大范围的结果责任。",
            ),
            "is_current": False,
        },
    ]

    alt_paths = []
    for item in (career_path or {}).get("alternative_paths") or []:
        alt_paths.append({"title": _safe_text(item.get("intermediate_role"), "相关岗位"), "skill_overlap": 65})
    for item in related_jobs:
        alt_paths.append(
            {
                "title": _safe_text(item.get("role_name"), "相关岗位"),
                "skill_overlap": _normalize_score(item.get("skill_overlap")),
            }
        )
    seen = set()
    deduped = []
    for item in alt_paths:
        title = item["title"]
        if title in seen:
            continue
        seen.add(title)
        deduped.append(item)
    return {"primary_path": primary_path, "alt_paths": deduped[:3]}


def _build_summary(student_profile: dict[str, Any], job_info: dict[str, Any], actions: list[dict[str, Any]]) -> str:
    basic = student_profile.get("basic_info") or {}
    background = _join_non_empty([_safe_text(basic.get("school")), _safe_text(basic.get("major"))], " / ")
    next_step = actions[0]["action"] if actions else "继续补强项目与实习证据"
    return _clean_paragraph(
        f"{_safe_text(basic.get('name'), '该学生')}当前与“{job_info['role_name']}”的综合匹配度为 {job_info['match_score']} 分。"
        f"{f'当前背景为 {background}。' if background else ''}"
        f"下一步建议优先执行：{next_step}"
    )


def _build_recommendations(job_info: dict[str, Any], dimensions: list[dict[str, Any]], actions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    score = job_info["match_score"]
    if score >= 80:
        primary = {
            "type": "positive",
            "title": "匹配基础较好",
            "content": f"你与“{job_info['role_name']}”已具备较好的进入门槛，接下来重点是强化项目证据和投递表达。",
        }
    elif score >= 60:
        primary = {
            "type": "improvement",
            "title": "适合集中补短板",
            "content": f"你与“{job_info['role_name']}”存在明确提升空间，补齐核心差距后有望明显抬升匹配度。",
        }
    else:
        primary = {
            "type": "warning",
            "title": "先补核心能力再投递",
            "content": f"当前与“{job_info['role_name']}”仍有较大差距，更适合先做阶段性准备后再进入正式投递。",
        }
    weakest = sorted(dimensions, key=lambda item: item["score"])[:1]
    recommendations = [primary]
    if weakest:
        recommendations.append(
            {
                "type": "action",
                "title": f"优先提升{weakest[0]['label']}",
                "content": weakest[0]["reason"],
            }
        )
    if actions:
        recommendations.append(
            {
                "type": "action",
                "title": "先完成最高优先级动作",
                "content": actions[0]["action"],
            }
        )
    return recommendations


def _build_report_content(
    student_profile: dict[str, Any],
    matching_results: list[dict[str, Any]],
    career_path: dict[str, Any] | None,
) -> tuple[dict[str, Any], str, list[dict[str, Any]]]:
    top_match = matching_results[0]
    job_info = _extract_job_info(top_match)
    dimensions = _extract_dimensions(top_match)
    actions = _build_actions(top_match)
    action_groups = _group_action_items(actions)
    review_checkpoints = _build_review_checkpoints(actions, job_info["match_score"])
    related_jobs = [
        {
            "role_name": _extract_job_info(item)["role_name"],
            "skill_overlap": max(0.4, min(0.95, _extract_job_info(item)["match_score"] / 100)),
        }
        for item in matching_results[1:4]
    ]
    paths = _build_paths(student_profile, job_info["role_name"], career_path, related_jobs, actions)
    basic = student_profile.get("basic_info") or {}
    skills = _extract_skills(student_profile, 5)
    projects = _extract_projects(student_profile, 2)
    internships = _extract_internships(student_profile, 2)
    competitiveness = _normalize_score(student_profile.get("competitiveness_score"))
    strongest_dimension = max(dimensions, key=lambda item: item["score"], default=None)
    weakest_dimension = min(dimensions, key=lambda item: item["score"], default=None)
    score_band = _score_band_text(job_info["match_score"])
    chapter_one = _clean_paragraph(
        f"{_safe_text(basic.get('name'), '该学生')}目前的综合竞争力约为 {competitiveness} 分，"
        f"{f'当前背景为 {_join_non_empty([_safe_text(basic.get('school')), _safe_text(basic.get('major'))], ' / ')}，' if _join_non_empty([_safe_text(basic.get('school')), _safe_text(basic.get('major'))], ' / ') else ''}"
        f"面向“{job_info['role_name']}”方向整体{score_band}。"
        f"{f'核心技能集中在 {", ".join(skills)}。' if skills else ''}"
        f"{f'已有 {", ".join(internships)} 等经历。' if internships else ''}"
        f"{f'项目实践包括 {", ".join(projects)}。' if projects else ''}"
    )
    chapter_two = _clean_paragraph(
        f"目标岗位为“{job_info['role_name']}”，当前综合匹配度为 {job_info['match_score']} 分。"
        f"{f'当前优势主要体现在“{strongest_dimension['label']}”维度。' if strongest_dimension else ''}"
        f"{f'短板则集中在“{weakest_dimension['label']}”维度，需要优先补强。' if weakest_dimension else ''}"
    )
    must_fix = len([item for item in actions if item["priority"] == "必须补齐"])
    recoverable_score = min(
        100,
        job_info["match_score"] + sum(abs(min(0, int(item.get("score_impact") or 0))) for item in actions[:3]),
    )
    chapter_three = _clean_paragraph(
        f"围绕“{job_info['role_name']}”的进入门槛，当前共识别出 {len(actions)} 项关键动作，其中必须补齐 {must_fix} 项。"
        f"优先级越高的事项越应尽快转化为项目、实习或简历中的可验证证据，若核心短板补齐，综合分预计可提升到 {recoverable_score} 分左右。"
    )
    chapter_four = _clean_paragraph(
        f"推荐以“{job_info['role_name']}”作为当前主路径，先进入准备期，再向初级岗位和进阶岗位逐步推进。"
        f"{f'同时也可关注 {", ".join(item["title"] for item in paths["alt_paths"])} 等相邻岗位作为横向备选。' if paths['alt_paths'] else ''}"
    )
    checkpoints = "、".join(item["item"] for item in actions[:3]) or "核心技能与项目表达"
    chapter_five = _clean_paragraph(
        f"建议至少每 3 个月做一次结构化复盘，重点检查 {checkpoints} 是否已经形成可展示成果，并同步观察简历表达是否更贴近岗位要求。"
        f"在第 6 个月节点重新上传最新简历并复跑匹配与报告，可以直观看到分数与路径建议的变化。"
    )
    summary = _build_summary(student_profile, job_info, actions)
    chapters = [
        {"chapter_id": 1, "title": _chapter_title(1), "text": chapter_one, "data": None, "status": "done"},
        {"chapter_id": 2, "title": _chapter_title(2), "text": chapter_two, "data": {"overall_score": job_info["match_score"], "dimensions": dimensions}, "status": "done"},
        {"chapter_id": 3, "title": _chapter_title(3), "text": chapter_three, "data": action_groups, "status": "done"},
        {"chapter_id": 4, "title": _chapter_title(4), "text": chapter_four, "data": paths, "status": "done"},
        {"chapter_id": 5, "title": _chapter_title(5), "text": chapter_five, "data": review_checkpoints, "status": "done"},
    ]
    recommendations = _build_recommendations(job_info, dimensions, actions)
    content_json = {
        "title": f"{_safe_text(basic.get('name'), '学生')} - {job_info['role_name']}职业发展报告",
        "summary": summary,
        "target_job": job_info,
        "dimensions": dimensions,
        "actions": actions,
        "action_groups": action_groups,
        "paths": paths,
        "review_checkpoints": review_checkpoints,
        "chapters": chapters,
        "metadata": {"generated_at": datetime.now(UTC).isoformat(), "template_version": REPORT_TEMPLATE_VERSION},
    }
    return content_json, summary, recommendations


async def _load_student_profile(student_id: UUID, db: AsyncSession) -> tuple[Student, dict[str, Any]]:
    student = await db.get(Student, student_id)
    if not student:
        raise ValueError("Student not found")
    result = await db.execute(select(StudentProfile).where(StudentProfile.student_id == student_id))
    profile = result.scalar_one_or_none()
    if profile is None or not profile.profile_json:
        raise ValueError("学生画像不存在，请先生成学生画像")
    return student, normalize_student_profile_json(profile.profile_json or {}, student)


def _match_to_payload(match_result: Any) -> dict[str, Any]:
    return {
        "job_id": str(match_result.job_profile_id),
        "job_profile_id": str(match_result.job_profile_id),
        "total_score": match_result.total_score,
        "scores_json": dict(match_result.scores_json or {}),
        "gaps_json": list(match_result.gaps_json or []),
    }


async def _resolve_matching_results(db: AsyncSession, student_id: UUID, target_job_ids: list[UUID] | None = None) -> list[dict[str, Any]]:
    if target_job_ids:
        results = []
        for job_id in target_job_ids:
            results.append(_match_to_payload(await match_student_job(db, student_id, job_id, mode="deep")))
        if not results:
            raise ValueError("未能生成目标岗位的匹配结果")
        return results
    recommended = await recommend_jobs(db, student_id, top_k=5)
    if not recommended:
        raise ValueError("暂无岗位匹配结果，请先完成匹配推荐")
    return [_match_to_payload(item) for item in recommended]


async def _resolve_career_path(db: AsyncSession, student_profile: dict[str, Any], matching_results: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not matching_results:
        return None
    target_role = _extract_job_info(matching_results[0])["role_name"]
    try:
        return await find_path_with_student_profile(db, student_profile, target_role, "expert")
    except Exception as exc:
        logger.warning("Career path generation failed for %s: %s", target_role, exc)
        return None


async def generate_outline(
    student_profile: dict[str, Any],
    matching_results: list[dict[str, Any]],
    career_path: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """生成报告大纲（目前使用模板大纲，后续可扩展为 LLM 生成）"""
    return {
        "title": DEFAULT_REPORT_TITLE,
        "chapters": REPORT_CHAPTERS,
        "generated_by": "template",
    }


def _extract_llm_text_and_data(raw_text: str) -> tuple[str, Any | None]:
    text = str(raw_text or "").strip()
    json_data: Any | None = None
    json_match = re.search(r"```json\s*([\s\S]*?)```", text, flags=re.IGNORECASE)
    if json_match:
        try:
            json_data = json.loads(json_match.group(1).strip())
            text = f"{text[:json_match.start()].strip()} {text[json_match.end():].strip()}".strip()
        except json.JSONDecodeError as exc:
            logger.warning("Failed to parse report chapter JSON block: %s", exc)
    return (_clean_paragraph(text) if text else ""), json_data


def _normalize_llm_dimensions(
    payload: Any,
    fallback_dimensions: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        normalized: list[dict[str, Any]] = []
        for item in payload:
            if not isinstance(item, dict):
                continue
            key = _safe_text(item.get("key"), "unknown")
            normalized.append(
                {
                    "key": key,
                    "label": _safe_text(item.get("label"), key),
                    "score": _normalize_score(item.get("score")),
                    "reason": _safe_text(item.get("reason"), "可继续补强这一维度的证据表达。"),
                }
            )
        return normalized or fallback_dimensions

    if not isinstance(payload, dict):
        return fallback_dimensions

    fallback_by_key = {
        _safe_text(item.get("key")): item
        for item in fallback_dimensions
        if isinstance(item, dict)
    }
    normalized = []
    for key, label in DIMENSION_META:
        current = payload.get(key)
        if key == "skill" and current is None:
            current = payload.get("skills")
        fallback = fallback_by_key.get(key) or {}
        current_dict = current if isinstance(current, dict) else {}
        normalized.append(
            {
                "key": key,
                "label": label,
                "score": _normalize_score(
                    current_dict.get("score")
                    if current_dict
                    else fallback.get("score")
                ),
                "reason": _safe_text(
                    current_dict.get("reason")
                    if current_dict
                    else fallback.get("reason"),
                    _safe_text(fallback.get("reason"), "可继续补强这一维度的证据表达。"),
                ),
            }
        )
    return normalized


def _normalize_llm_chapter_data(
    chapter_id: int,
    json_data: Any,
    fallback_data: Any,
    fallback_dimensions: list[dict[str, Any]],
    actions: list[dict[str, Any]],
    baseline_score: int,
    target_role: str,
) -> Any:
    if json_data is None:
        return fallback_data

    if chapter_id == 2 and isinstance(json_data, dict):
        return {
            "overall_score": _normalize_score(json_data.get("overall_score") or baseline_score),
            "dimensions": _normalize_llm_dimensions(
                json_data.get("dimensions"),
                fallback_dimensions,
            ),
        }

    if chapter_id == 3:
        return _group_action_items(_normalize_action_items(json_data))

    if chapter_id == 4 and isinstance(json_data, dict):
        return _normalize_paths(json_data, target_role, actions)

    if chapter_id == 5:
        return _normalize_review_checkpoints(json_data, actions, baseline_score)

    return fallback_data


async def generate_chapters(
    outline: dict[str, Any],
    student_profile: dict[str, Any],
    matching_results: list[dict[str, Any]],
    career_path: dict[str, Any] | None,
    db: AsyncSession,
) -> list[dict[str, Any]]:
    """逐章调用 LLM 生成报告正文，失败时回退到模板内容。"""
    del outline, db

    content_json, _, _ = _build_report_content(student_profile, matching_results, career_path)
    template_chapters: list[dict[str, Any]] = list(content_json.get("chapters") or [])
    template_dimensions = list(content_json.get("dimensions") or [])
    actions = _normalize_action_items(content_json.get("actions") or [])

    basic_info = student_profile.get("basic_info") or {}
    name = _safe_text(basic_info.get("name"), "同学")
    school_major = _join_non_empty(
        [
            _safe_text(basic_info.get("school")),
            _safe_text(basic_info.get("major")),
        ]
    )
    overall_score = _normalize_score(
        student_profile.get("competitiveness_score")
        or student_profile.get("overall_score")
    )
    percentile = max(1, min(99, 100 - overall_score))

    top_skills = ", ".join(_extract_skills(student_profile, limit=5)) or "暂无技能数据"
    internships_text = "；".join(_extract_internships(student_profile, limit=2)) or "暂无实习经历"
    projects_text = "；".join(_extract_projects(student_profile, limit=2)) or "暂无项目经历"
    certs_text = ", ".join(_extract_certificates(student_profile, limit=3)) or "暂无"

    soft_skills_raw = student_profile.get("soft_competencies") or {}
    soft_skills_parts = []
    if isinstance(soft_skills_raw, dict):
        for label, payload in soft_skills_raw.items():
            if isinstance(payload, dict):
                score = payload.get("value") or payload.get("score") or 0
                if score:
                    soft_skills_parts.append(f"{label}({score}/5)")
    soft_skills_text = "、".join(soft_skills_parts[:4]) or "暂无软素养数据"

    top_match = matching_results[0] if matching_results else {}
    job_info = _extract_job_info(top_match)
    target_job_name = job_info.get("role_name") or "目标岗位"
    overall_match_score = _normalize_score(
        top_match.get("total_score") or job_info.get("match_score")
    )

    scores_json = dict(top_match.get("scores_json") or {})
    basic_score = _normalize_score((scores_json.get("basic") or {}).get("score"))
    skills_score = _normalize_score((scores_json.get("skill") or {}).get("score"))
    competency_score = _normalize_score((scores_json.get("competency") or {}).get("score"))
    potential_score = _normalize_score((scores_json.get("potential") or {}).get("score"))

    gaps_json = list(top_match.get("gaps_json") or [])
    main_gaps_text = "；".join(
        g.get("gap_item") or g.get("item") or ""
        for g in gaps_json[:4]
        if isinstance(g, dict)
    ) or "暂无差距数据"
    match_highlights = "；".join(
        r for r in (scores_json.get("match_reasons") or [])[:3]
        if isinstance(r, str)
    ) or "综合能力较为均衡"

    gap_list_text = "\n".join(
        f"- {g.get('gap_item', '')}: {g.get('suggestion', g.get('gap_desc', ''))}"
        for g in gaps_json[:6]
        if isinstance(g, dict)
    ) or "暂无具体差距"
    student_skills_text = top_skills

    paths_data = content_json.get("paths") or {}
    related_jobs_text = "、".join(
        _safe_text(p.get("title") or p.get("role_name"))
        for p in (paths_data.get("alt_paths") or [])[:3]
        if isinstance(p, dict) and _safe_text(p.get("title") or p.get("role_name"))
    ) or "同类相关岗位"
    student_summary = f"{name}，{school_major or '专业背景待补充'}，竞争力评分 {overall_score}/100"

    actions_template = _group_action_items(actions).get("short_term") or []
    action_summary = "；".join(
        a.get("item") or "" for a in actions_template[:3] if isinstance(a, dict)
    ) or main_gaps_text

    chapter_prompts = [
        CHAPTER_1_PROMPT.format(
            name=name,
            school=school_major or "学校与专业待补充",
            overall_score=overall_score,
            percentile=percentile,
            top_skills=top_skills,
            internships=internships_text,
            projects=projects_text,
            certificates=certs_text,
            soft_skills=soft_skills_text,
            target_job_name=target_job_name,
        ),
        CHAPTER_2_PROMPT.format(
            target_job_name=target_job_name,
            overall_score=overall_match_score,
            basic_score=basic_score,
            skills_score=skills_score,
            competency_score=competency_score,
            potential_score=potential_score,
            match_highlights=match_highlights,
            main_gaps=main_gaps_text,
        ),
        CHAPTER_3_PROMPT.format(
            gap_list=gap_list_text,
            student_skills=student_skills_text,
        ),
        CHAPTER_4_PROMPT.format(
            student_summary=student_summary,
            target_job_name=target_job_name,
            related_jobs=related_jobs_text,
            student_skills=student_skills_text,
        ),
        CHAPTER_5_PROMPT.format(
            action_summary=action_summary,
            main_gaps=main_gaps_text,
        ),
    ]

    final_chapters: list[dict[str, Any]] = []

    for i, prompt_text in enumerate(chapter_prompts):
        chapter_id = i + 1
        template_ch = next(
            (
                ch
                for ch in template_chapters
                if ch.get("chapter_id") == chapter_id
            ),
            template_chapters[i] if i < len(template_chapters) else {},
        )

        try:
            llm_text = await llm.generate(
                prompt=prompt_text,
                system_prompt=REPORT_SYSTEM_PROMPT,
                temperature=0.4,
                max_tokens=800,
                disable_reasoning=True,
                max_retries=1,
                timeout=12,
            )
            llm_text, json_data = _extract_llm_text_and_data(llm_text)
            chapter_data = _normalize_llm_chapter_data(
                chapter_id=chapter_id,
                json_data=json_data,
                fallback_data=template_ch.get("data"),
                fallback_dimensions=template_dimensions,
                actions=actions,
                baseline_score=overall_match_score,
                target_role=target_job_name,
            )

            final_chapters.append(
                {
                    **template_ch,
                    "chapter_id": chapter_id,
                    "title": _chapter_title(chapter_id),
                    "text": llm_text or template_ch.get("text") or "",
                    "data": chapter_data,
                    "status": "done",
                }
            )
            logger.info(
                "Report chapter %d generated by LLM (%d chars)",
                chapter_id,
                len(llm_text),
            )

        except Exception as exc:
            logger.warning("LLM chapter %d failed, using template: %s", chapter_id, exc)
            final_chapters.append(template_ch)

    return final_chapters


def _next_version(version: str | None) -> str:
    try:
        return f"{float(version or '1.0') + 0.1:.1f}"
    except (TypeError, ValueError):
        return "1.1"


async def merge_and_save(student_id: UUID, report_data: dict[str, Any], db: AsyncSession) -> CareerReport:
    content_json = normalize_report_content(dict(report_data.get("content_json") or {}))
    summary = _clean_paragraph(str(report_data.get("summary") or content_json.get("summary") or ""))
    recommendations = list(report_data.get("recommendations") or [])
    existing = (
        await db.execute(
            select(CareerReport)
            .where(CareerReport.student_id == student_id)
            .order_by(desc(CareerReport.updated_at), desc(CareerReport.created_at))
            .limit(1)
        )
    ).scalars().first()
    if existing is None:
        report = CareerReport(
            student_id=student_id,
            status="completed",
            version="1.0",
            content_json=content_json,
            summary=summary,
            recommendations=recommendations,
        )
        db.add(report)
        await db.flush()
    else:
        report = existing
        report.status = "completed"
        report.version = _next_version(report.version)
        report.content_json = content_json
        report.summary = summary
        report.recommendations = recommendations
        report.pdf_path = None
        report.docx_path = None
        await db.flush()
    db.add(
        ReportVersion(
            report_id=report.id,
            version=report.version,
            content=content_json,
            change_notes="重新生成报告",
        )
    )
    await db.commit()
    await db.refresh(report)
    return report


async def generate_full_report(student_id: UUID, db: AsyncSession, target_job_ids: list[UUID] | None = None) -> CareerReport:
    _, student_profile = await _load_student_profile(student_id, db)
    matching_results = await _resolve_matching_results(db, student_id, target_job_ids)
    career_path = await _resolve_career_path(db, student_profile, matching_results)
    outline = await generate_outline(student_profile, matching_results, career_path)
    chapters = await generate_chapters(outline, student_profile, matching_results, career_path, db)
    content_json, summary, recommendations = _build_report_content(student_profile, matching_results, career_path)
    content_json["chapters"] = chapters
    content_json["metadata"]["student_id"] = str(student_id)
    content_json["metadata"]["target_job_ids"] = [str(item) for item in target_job_ids or []]
    return await merge_and_save(
        student_id,
        {
            "content_json": content_json,
            "summary": summary,
            "recommendations": recommendations,
            "matching_results": matching_results,
            "career_path": career_path,
        },
        db,
    )


async def generate_report(student_id: UUID, db: AsyncSession, job_ids: list[UUID] | None = None) -> dict[str, Any]:
    return serialize_career_report(await generate_full_report(student_id, db, job_ids))


def _build_export_chart_html(content: dict[str, Any]) -> str:
    normalized = normalize_report_content(content)
    dimensions = normalized.get("dimensions") or []
    if not dimensions:
        matching_results = content.get("matching_results") or []
        if isinstance(matching_results, list) and matching_results:
            top_match = matching_results[0] if isinstance(matching_results[0], dict) else {}
            if isinstance(top_match, dict):
                dimensions = _extract_dimensions(top_match)
    if not dimensions:
        return ""
    rows = []
    for item in dimensions:
        rows.append(
            f'<div class="chart-row"><div class="chart-label">{_escape_html(item.get("label"))}</div>'
            f'<div class="chart-bar"><div class="chart-fill" style="width: {item.get("score", 0)}%"></div></div>'
            f'<div class="chart-value">{item.get("score", 0)}</div></div>'
        )
    return '<div class="chart-container"><h3>四维匹配概览</h3>' + "".join(rows) + "</div>"


def _build_export_html(report: CareerReport, content_json: dict[str, Any]) -> str:
    content = normalize_report_content(content_json)
    cards = []
    for chapter in content["chapters"]:
        extra_html = ""
        if chapter["chapter_id"] == 2:
            extra_html = _build_export_chart_html(content)
        elif chapter["chapter_id"] == 3 and isinstance(chapter.get("data"), dict):
            groups = []
            for title, key in (("短期 0-6 月", "short_term"), ("中期 6-18 月", "medium_term")):
                items = []
                for item in chapter["data"].get(key) or []:
                    resources = " / ".join(_escape_html(resource) for resource in (item.get("resources") or []))
                    items.append(
                        f'<div class="action-item"><strong>{_escape_html(item.get("priority"))} / {_escape_html(item.get("item"))}</strong>'
                        f'<div>{_escape_html(item.get("gap_desc"))}</div>'
                        f'<div>{_escape_html(item.get("action"))}</div>'
                        f'<div class="muted">周期：{_escape_html(item.get("timeline"))}'
                        f'{f" · 资源：{resources}" if resources else ""}</div></div>'
                    )
                if items:
                    groups.append(f'<div class="action-group-title">{title}</div>' + "".join(items))
            extra_html = '<div class="action-list">' + "".join(groups) + "</div>"
        elif chapter["chapter_id"] == 4 and isinstance(chapter.get("data"), dict):
            nodes = []
            for item in chapter["data"].get("primary_path") or []:
                nodes.append(
                    f'<div class="path-node"><div class="muted">{_escape_html(item.get("stage"))}</div>'
                    f'<div><strong>{_escape_html(item.get("title"))}</strong></div><div>{_escape_html(item.get("condition"))}</div></div>'
                )
            alt_titles = [f"<li>{_escape_html(item.get('title'))}</li>" for item in (chapter["data"].get("alt_paths") or [])]
            extra_html = '<div class="path-list">' + "".join(nodes) + "</div>" + (f"<ul>{''.join(alt_titles)}</ul>" if alt_titles else "")
        elif chapter["chapter_id"] == 5 and isinstance(chapter.get("data"), dict):
            checkpoints = []
            for item in chapter["data"].get("review_checkpoints") or []:
                checkpoints.append(
                    f'<div class="timeline-item"><strong>{_escape_html(item.get("month"))} 月</strong>'
                    f'<div>{_escape_html(item.get("goal"))}</div>'
                    f'<div class="muted">KPI：{_escape_html(item.get("kpi"))}</div>'
                    f'<div class="muted">验证动作：{_escape_html(item.get("action"))}</div></div>'
                )
            extra_html = '<div class="timeline-list">' + "".join(checkpoints) + "</div>"
        cards.append(
            f'<section class="card"><div class="card-head"><h2>{_escape_html(chapter["title"])}</h2><span class="badge">已生成</span></div>'
            f'<p>{_escape_html(chapter.get("text"))}</p>{extra_html}</section>'
        )
    rec_html = "".join(
        f'<div class="recommendation"><strong>{_escape_html(item.get("title"))}</strong><span>{_escape_html(item.get("content"))}</span></div>'
        for item in (report.recommendations or [])
    )
    title = content.get("title") or DEFAULT_REPORT_TITLE
    summary = report.summary or content.get("summary") or ""
    generated_at = report.created_at.strftime("%Y-%m-%d %H:%M:%S") if report.created_at else datetime.now(UTC).strftime("%Y-%m-%d %H:%M:%S")
    return f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="UTF-8"><title>{_escape_html(title)}</title>
<style>
body{{font-family:"Microsoft YaHei","PingFang SC",sans-serif;background:#f4f6fb;color:#111827;margin:0;padding:32px}}
.container{{max-width:860px;margin:0 auto}} .hero,.card,.panel{{background:#fff;border:1px solid #e5e7eb;border-radius:20px;padding:24px;margin-bottom:18px;box-shadow:0 10px 30px rgba(15,23,42,.06)}}
.hero{{background:linear-gradient(135deg,#4f46e5,#2563eb);color:#fff}} .hero h1{{margin:0 0 10px}} .hero p{{margin:0;line-height:1.8}}
.card-head{{display:flex;justify-content:space-between;align-items:center;gap:12px;margin-bottom:12px}} .card-head h2{{margin:0;font-size:22px}}
.badge{{background:#dcfce7;color:#166534;border-radius:999px;padding:6px 12px;font-size:12px;font-weight:700}} p{{line-height:1.9;white-space:pre-wrap}}
.chart-container,.action-item,.path-node,.recommendation,.timeline-item{{background:#f8fafc;border-radius:16px;padding:14px 16px;margin-top:12px}} .muted{{color:#6b7280;font-size:12px}}
.chart-row{{display:flex;align-items:center;gap:12px;margin:10px 0}} .chart-label{{width:72px;font-weight:700}} .chart-bar{{flex:1;height:10px;background:#e5e7eb;border-radius:999px;overflow:hidden}} .chart-fill{{height:100%;background:linear-gradient(90deg,#60a5fa,#2563eb)}} .chart-value{{width:40px;text-align:right}}
.action-group-title{{margin-top:16px;font-size:14px;font-weight:700;color:#1f2937}} .timeline-list{{display:grid;gap:12px}}
.recommendation{{display:grid;gap:6px}} .footer{{text-align:center;color:#6b7280;font-size:12px;margin-top:18px}}
</style></head><body><div class="container"><section class="hero"><h1>{_escape_html(title)}</h1><p>{_escape_html(summary)}</p></section>{''.join(cards)}<section class="panel"><h3>推荐建议</h3>{rec_html or '<div class="recommendation">当前暂无额外建议。</div>'}</section><div class="footer"><div>生成时间：{_escape_html(generated_at)}</div><div>版本：{_escape_html(report.version or '1.0')}</div></div></div></body></html>"""


async def export_to_pdf(report_id: UUID, db: AsyncSession) -> str:
    report = await db.get(CareerReport, report_id)
    if report is None:
        raise ValueError(f"Report {report_id} not found")
    html_content = _build_export_html(report, report.content_json or {})
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(UTC).strftime("%Y%m%d%H%M%S")
    html_path = EXPORT_DIR / f"career_report_{report.student_id}_{timestamp}.html"
    pdf_path = EXPORT_DIR / f"career_report_{report.student_id}_{timestamp}.pdf"
    html_path.write_text(html_content, encoding="utf-8")
    script = """
import asyncio
import sys
from pathlib import Path
from playwright.async_api import async_playwright

async def main():
    html_path = Path(sys.argv[1]).resolve()
    pdf_path = Path(sys.argv[2]).resolve()
    async with async_playwright() as playwright:
        browser = await playwright.chromium.launch()
        page = await browser.new_page()
        await page.goto(html_path.as_uri(), wait_until='networkidle')
        await page.pdf(path=str(pdf_path), format='A4', print_background=True, margin={'top': '18mm', 'right': '14mm', 'bottom': '18mm', 'left': '14mm'})
        await browser.close()

asyncio.run(main())
"""
    try:
        result = await asyncio.to_thread(
            subprocess.run,
            [sys.executable, "-c", script, str(html_path), str(pdf_path)],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError((result.stderr or result.stdout or "").strip())
    except Exception as exc:  # pragma: no cover
        logger.error("PDF export failed for report %s: %s", report_id, exc)
        raise RuntimeError(f"PDF 导出失败: {exc}") from exc
    report.pdf_path = str(pdf_path)
    await db.commit()
    return str(pdf_path)


async def _export_to_html(report_id: UUID, db: AsyncSession) -> str:
    report = await db.get(CareerReport, report_id)
    if report is None:
        raise ValueError(f"Report {report_id} not found")
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    html_path = EXPORT_DIR / f"career_report_{report.student_id}_{datetime.now(UTC).strftime('%Y%m%d%H%M%S')}.html"
    html_path.write_text(_build_export_html(report, report.content_json or {}), encoding="utf-8")
    return str(html_path)


async def export_to_docx(report_id: UUID, db: AsyncSession) -> str:
    report = await db.get(CareerReport, report_id)
    if report is None:
        raise ValueError(f"Report {report_id} not found")
    from docx import Document
    content = normalize_report_content(report.content_json or {})
    document = Document()
    document.add_heading(content.get("title") or DEFAULT_REPORT_TITLE, 0)
    if report.summary:
        document.add_heading("报告摘要", level=1)
        document.add_paragraph(report.summary)
    for chapter in content["chapters"]:
        document.add_heading(chapter["title"], level=1)
        if chapter.get("text"):
            document.add_paragraph(chapter["text"])
        if chapter["chapter_id"] == 3 and isinstance(chapter.get("data"), dict):
            for title, key in (("短期 0-6 月", "short_term"), ("中期 6-18 月", "medium_term")):
                items = chapter["data"].get(key) or []
                if not items:
                    continue
                document.add_paragraph(title)
                for item in items:
                    document.add_paragraph(
                        f"{item.get('priority')} - {item.get('item')}: {item.get('action')}",
                        style="List Bullet",
                    )
        if chapter["chapter_id"] == 5 and isinstance(chapter.get("data"), dict):
            for item in chapter["data"].get("review_checkpoints") or []:
                document.add_paragraph(
                    f"{item.get('month')}月 - {item.get('goal')} / KPI: {item.get('kpi')}",
                    style="List Bullet",
                )
    if report.recommendations:
        document.add_heading("推荐建议", level=1)
        for item in report.recommendations:
            document.add_paragraph(f"{item.get('title')}: {item.get('content')}", style="List Bullet")
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = EXPORT_DIR / f"career_report_{report.student_id}_{datetime.now(UTC).strftime('%Y%m%d%H%M%S')}.docx"
    document.save(str(docx_path))
    report.docx_path = str(docx_path)
    await db.commit()
    return str(docx_path)


def _build_report_polish_context(content: dict[str, Any], report: CareerReport) -> dict[str, Any]:
    return {
        "title": content.get("title") or DEFAULT_REPORT_TITLE,
        "summary": content.get("summary") or report.summary or "",
        "target_job": content.get("target_job") or {},
        "dimensions": content.get("dimensions") or [],
        "actions": content.get("actions") or [],
        "review_checkpoints": content.get("review_checkpoints") or {},
        "paths": content.get("paths") or {},
        "chapters": [
            {
                "chapter_id": chapter.get("chapter_id"),
                "title": chapter.get("title"),
                "text": chapter.get("text") or "",
            }
            for chapter in (content.get("chapters") or [])
        ],
    }


async def _ai_polish_report_content(content: dict[str, Any], report: CareerReport) -> tuple[dict[str, Any], list[str]]:
    context = _build_report_polish_context(content, report)
    prompt = REPORT_POLISH_USER_TEMPLATE.format(
        report_context=json.dumps(context, ensure_ascii=False, indent=2)
    )
    result = await llm.generate_json(
        prompt=prompt,
        system_prompt=REPORT_POLISH_SYSTEM_PROMPT,
        temperature=0.4,
        max_retries=2,
        timeout=25,
    )

    polished_summary = _clean_paragraph(_safe_text(result.get("summary")))
    raw_chapters = result.get("chapters")
    if not isinstance(raw_chapters, list):
        raise ValueError("AI 润色结果缺少 chapters 数组")

    current_chapters = {chapter["chapter_id"]: dict(chapter) for chapter in content.get("chapters") or []}
    changes: list[str] = []
    polished_chapters: list[dict[str, Any]] = []
    seen_ids: set[int] = set()

    for item in raw_chapters:
        if not isinstance(item, dict):
            continue
        chapter_id = int(item.get("chapter_id") or 0)
        original = current_chapters.get(chapter_id)
        if original is None:
            continue
        polished_text = _clean_paragraph(_safe_text(item.get("text")))
        if not polished_text:
            raise ValueError(f"AI 润色结果缺少第 {chapter_id} 章正文")
        if polished_text != (original.get("text") or ""):
            changes.append(f"AI 改写了《{original['title']}》正文")
        original["text"] = polished_text
        polished_chapters.append(original)
        seen_ids.add(chapter_id)

    missing_ids = sorted(set(current_chapters) - seen_ids)
    if missing_ids:
        raise ValueError(f"AI 润色结果缺少章节: {', '.join(str(item) for item in missing_ids)}")

    polished_chapters.sort(key=lambda item: item["chapter_id"])
    updated_content = dict(content)
    if polished_summary and polished_summary != (content.get("summary") or ""):
        changes.insert(0, "AI 改写了报告摘要")
    updated_content["summary"] = polished_summary or content.get("summary") or ""
    updated_content["chapters"] = polished_chapters
    updated_content["metadata"] = {
        **dict(content.get("metadata") or {}),
        "last_polished_at": datetime.now(UTC).isoformat(),
        "last_polish_mode": "ai",
    }
    return updated_content, changes


async def polish_report(report_id: UUID, db: AsyncSession) -> dict[str, Any]:
    report = await db.get(CareerReport, report_id)
    if report is None:
        return {"polished": False, "error": "Report not found"}
    content = normalize_report_content(report.content_json or {})
    if not any((chapter.get("text") or "").strip() for chapter in content.get("chapters") or []):
        return {"polished": False, "error": "报告正文为空，请先生成报告"}

    updated_content, changes = await _ai_polish_report_content(content, report)
    if not changes:
        return {"polished": True, "changes": [], "version": report.version or "1.0"}

    report.content_json = updated_content
    report.summary = updated_content.get("summary") or report.summary
    report.version = _next_version(report.version)
    db.add(
        ReportVersion(
            report_id=report.id,
            version=report.version,
            content=updated_content,
            change_notes="AI 增强润色报告文案",
        )
    )
    await db.commit()
    return {"polished": True, "changes": changes, "version": report.version}


async def check_completeness(report_id: UUID, db: AsyncSession) -> dict[str, Any]:
    report = await db.get(CareerReport, report_id)
    if report is None:
        raise ValueError("Report not found")
    content = normalize_report_content(report.content_json or {})
    missing = []
    if not content.get("summary"):
        missing.append("报告摘要")
    if not content.get("dimensions"):
        missing.append("四维匹配结果")
    if not content.get("actions"):
        missing.append("行动计划")
    if not (content.get("paths") or {}).get("primary_path"):
        missing.append("职业路径规划")
    for chapter in REPORT_CHAPTERS:
        current = next((item for item in content["chapters"] if item["chapter_id"] == chapter["chapter_id"]), None)
        if current is None or not current.get("text"):
            missing.append(chapter["title"])
    suggestions = []
    if "四维匹配结果" in missing:
        suggestions.append("先确保学生画像和岗位匹配结果已生成，再重新生成报告。")
    if "行动计划" in missing:
        suggestions.append("补充至少 3 条可执行行动项，并给出时间周期。")
    if "职业路径规划" in missing:
        suggestions.append("补充主路径和横向备选岗位，便于后续导出与复盘。")
    if not suggestions and missing:
        suggestions.append("重新生成一次报告以补齐缺失内容。")
    return {"complete": not missing, "missing_items": missing, "suggestions": suggestions, "chapter_count": len(content["chapters"])}


async def create_report_version(report_id: UUID, version: str, db: AsyncSession, change_notes: str | None = None) -> dict[str, Any]:
    report = await db.get(CareerReport, report_id)
    if report is None:
        raise ValueError(f"Report {report_id} not found")
    report_version = ReportVersion(report_id=report.id, version=version, content=normalize_report_content(report.content_json or {}), change_notes=change_notes)
    db.add(report_version)
    await db.commit()
    return {"id": str(report_version.id), "report_id": str(report_version.report_id), "version": report_version.version, "created_at": report_version.created_at.isoformat() if report_version.created_at else None}
